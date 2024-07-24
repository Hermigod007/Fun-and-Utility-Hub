import argparse
import os
import sys
from docx import Document
import re
import zlib
import base64
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import gzip
import shutil


class AddingDocument:
    def __init__(self):
        self.generated_file = ""

    def writefile(self):
        document = Document()
        heading = input("Please add heading to your Document: ")
        document.add_heading(heading, 0)

        path = os.path.expanduser("~/Desktop/")

        info = input("Please enter the information in your file: ")
        document.add_paragraph(info)

        self.generated_file = input("Save the file named as: ")
        self.generated_file = self.generated_file + ".docx"
        document.save(os.path.join(path, self.generated_file))

        doc = Document(os.path.join(path, self.generated_file))
        completed_text = [paras.text for paras in doc.paragraphs]

        return "\n".join(completed_text)

    def addfile(self, filename):
        doc = Document(filename)
        completed_text = [paras.text for paras in doc.paragraphs]
        return "\n".join(completed_text)


class Functionality:
    def text_manipulation(self, data):
        operation = int(
            input(
                "1) Search and Replace \n2) Text Extraction \n3) Text Sorting \n\nPlease kindly select the text manipulation operation which you want to perform: "
            )
        )

        if operation == 1:
            search, replace = input(
                "Please enter the word which you want to search and replace (separated by space): "
            ).split()
            replaced_text = re.sub(search, replace, data)
            print("\nThe changes have been made in the file: \n")
            print(replaced_text)

        elif operation == 2:
            word = input("Which word do you want to search for: ")
            count = data.lower().split().count(word.lower())
            total_words = len(data.split())

            print(
                f"The word '{word}' has appeared {count} times in the given data.")
            print(f"Total Word Count in Document is: {total_words}\n")

        elif operation == 3:
            words = data.split()
            words.sort()
            sorted_text = " ".join(words)
            print("The sorted text is:\n")
            print(sorted_text)

    def file_compression(self, data):
        compressed_data = base64.b64encode(
            zlib.compress(data.encode("utf-8"), 9)).decode("utf-8")
        with open(os.path.expanduser("~/Desktop/Compressed.txt"), "w") as compressed_file:
            compressed_file.write(compressed_data)
        print("The file has been compressed and saved as Compressed.txt")

    def create_pdf(self, docx_file):
        document = Document(docx_file)
        pdf_file = docx_file.replace(".docx", ".pdf")
        pdf = canvas.Canvas(pdf_file, pagesize=letter)
        pdf.setFont("Helvetica", 12)
        width, height = letter
        y = height - 40

        for paragraph in document.paragraphs:
            text = paragraph.text
            pdf.drawString(40, y, text)
            y -= 20
            if y < 40:
                pdf.showPage()
                y = height - 40

        pdf.save()
        print(f"PDF created successfully: {pdf_file}")


def count_words(file_path):
    doc = Document(file_path)
    words = []
    for paragraph in doc.paragraphs:
        words.extend(paragraph.text.split())
    return len(words)


def count_characters(file_path):
    doc = Document(file_path)
    text = "".join([paragraph.text for paragraph in doc.paragraphs])
    return len(text)


def count_lines(file_path):
    doc = Document(file_path)
    return len(doc.paragraphs)


def sort_lines_alphabetically(file_path):
    doc = Document(file_path)
    lines = [paragraph.text for paragraph in doc.paragraphs]
    sorted_lines = sorted(lines, key=str.lower)
    return "\n".join(sorted_lines)


def remove_duplicate_lines(file_path):
    doc = Document(file_path)
    lines = [paragraph.text for paragraph in doc.paragraphs]
    unique_lines = list(set(lines))
    unique_lines.sort(key=str.lower)
    return "\n".join(unique_lines)


def replace_word(file_path, target, replacement):
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        if target in paragraph.text:
            paragraph.text = paragraph.text.replace(target, replacement)
    doc.save(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])


def compress_file(file_path):
    with open(file_path, "rb") as f_in:
        with gzip.open(file_path + ".gz", "wb") as f_out:
            shutil.copyfileobj(f_in, f_out)
    return file_path + ".gz"


def decompress_file(file_path):
    with gzip.open(file_path, "rb") as f_in:
        with open(file_path[:-3], "wb") as f_out:
            shutil.copyfileobj(f_in, f_out)
    return file_path[:-3]


def cli_main():
    parser = argparse.ArgumentParser(description="Text manipulation utility.")
    parser.add_argument("file", type=str, help="Path to the DOCX file")
    parser.add_argument("-w", "--words", action="store_true",
                        help="Count words in the text file.")
    parser.add_argument("-c", "--characters", action="store_true",
                        help="Count characters in the text file.")
    parser.add_argument("-l", "--lines", action="store_true",
                        help="Count lines in the text file.")
    parser.add_argument("-s", "--sort", action="store_true",
                        help="Sort lines in the text file alphabetically.")
    parser.add_argument("-d", "--distinct", action="store_true",
                        help="Remove duplicate lines from the text file.")
    parser.add_argument("-r", "--replace", nargs=2, metavar=("TARGET",
                        "REPLACEMENT"), help="Replace a word in the text file.")
    parser.add_argument("-z", "--compress", action="store_true",
                        help="Compress the text file using gzip.")
    parser.add_argument("-x", "--decompress", action="store_true",
                        help="Decompress the gzip compressed text file.")
    parser.add_argument("-p", "--pdf", action="store_true",
                        help="Create a PDF from the DOCX file.")

    args = parser.parse_args()

    if args.words:
        print(f"Words: {count_words(args.file)}")
    elif args.characters:
        print(f"Characters: {count_characters(args.file)}")
    elif args.lines:
        print(f"Lines: {count_lines(args.file)}")
    elif args.sort:
        print(f"Sorted Lines:\n{sort_lines_alphabetically(args.file)}")
    elif args.distinct:
        print(f"Unique Lines:\n{remove_duplicate_lines(args.file)}")
    elif args.replace:
        target, replacement = args.replace
        print(f"Updated Text:\n{replace_word(args.file, target, replacement)}")
    elif args.compress:
        compressed_file = compress_file(args.file)
        print(f"File compressed successfully: {compressed_file}")
    elif args.decompress:
        decompressed_file = decompress_file(args.file)
        print(f"File decompressed successfully: {decompressed_file}")
    elif args.pdf:
        functionality = Functionality()
        functionality.create_pdf(args.file)
    else:
        print("No valid operation specified.", file=sys.stderr)
        sys.exit(1)


def interactive_main():
    print("\n.........Welcome to the Command line prompt File Manipulation Program.........\n")
    method = int(
        input("1) To Create a new File\n2) To Upload a File\nPlease Choose an Option: "))

    information = ''

    if method == 1:
        test_file = AddingDocument()
        information = test_file.writefile()
    elif method == 2:
        test_file = AddingDocument()
        path = input(
            "Please enter the Path/location of where the file is saved: ")
        information = test_file.addfile(path)
    else:
        print("Please Select a valid input")
        return

    raw_data = Functionality()

    while True:
        choice = int(input(
            "\n1) Text Manipulation \n2) File Compression \n 3) Exit \n\nWhat operation do you want to perform: "))
        print('\n')

        if choice == 1:
            raw_data.text_manipulation(information)
        elif choice == 2:
            raw_data.file_compression(information)
        elif choice == 3:
            raw_data.create_pdf(test_file.generated_file)
        elif choice == 4:
            break
        else:
            print("Please select a valid option.")


if __name__ == "__main__":

    if len(sys.argv) > 1:
        cli_main()
    else:

        print("To run the following functions in command line use this commands")
        print(" Command to count total number of words in text file : -w\n",
              "Command to count total number of characters in text file : -c \n",
              "Command to count total number of lines in text file : -l \n",
              "Command to sort the lines in alphabetical order in text file : -s \n",
              "Command to remove duplicate words from the text file : -d \n",
              "Command to replace the words in text file : -r \n",
              "Command to compress the text file to zip file : -z \n",
              "Command to decompress the zip file : -x \n",
              "Command to create a PDF file from Word or .docx file : -p \n\n",)
        interactive_main()
